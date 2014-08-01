import string
import os
import xml.etree.ElementTree as ET
from xml.dom import minidom
import sys
import time
from datetime import date
import re
from collections import OrderedDict
from operator import itemgetter

import codecs

from docx import Document

from lxml import etree

import zipfile


#Patient #1

def doc2text(dirIn1,dirOut):
        #create folder
    printHeadLine=True
    if os.path.exists(dirOut) == False:    
            os.mkdir(dirOut)
            
    #Explore Directories
    for dirname, dirnames, filenames in os.walk(dirIn):
        for filename in filenames:
            if filename.strip()[0]=='.' :
                continue
            f = os.path.join(dirname, filename) 
         
           # print "f: ",f
            x=''
            file=open(f)
            document=Document(file)
            paragphs=document.paragraphs
            for aPara in paragphs:
                x+=aPara.text
                x+='\n'
            file.close()
            
            outName=re.split('\.',filename)[0]+'.txt'
            outf=os.path.join(dirOut,outName)
            outFile=codecs.open(outf,'w','utf-8')
            outFile.write(x)
            outFile.close()
            print outName
            
            #     x=''
#     fName='kd_data/Patient HPIs for review.docx'
#     f=open(fName)
#     document = Document(f)
#     pgraphs=document.paragraphs
#     for aPara in pgraphs:
#         x+=aPara.text
#         x+='\n'
#     print x.encode('utf-8')   
#     f.close()

            
            
    print 'files created'

    
if __name__=="__main__":
#      dirIn = sys.argv[1]
#      dirIn2= sys.argv[2]
#      dirOut=sys.argv[3]
#      doc2text(dirIn,dirOut)

    dirIn = 'kd_data'
    dirOut= 'kd_data_text'
    doc2text(dirIn,dirOut)

    
#     x=''
#     fName='kd_data/Patient HPIs for review.docx'
#     f=open(fName)
#     document = Document(f)
#     pgraphs=document.paragraphs
#     for aPara in pgraphs:
#         x+=aPara.text
#         x+='\n'
#     print x.encode('utf-8')   
#     f.close()

#       docx=zipfile.ZipFile('kd_data/Patient HPIs for review.docx')
#       content=docx.read('word/document.xml')
#       cleaned=re.sub('<(.|\n)*?>','',content)
# #       print cleaned.encode('utf-8')
#       print cleaned

#     x=''
#     fName='kd_data/Patient HPIs for review.docx'
#     document = opendocx(fName)
#     print getdocumenttext(document).encode('utf-8')


